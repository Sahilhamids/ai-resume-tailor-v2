from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx_from_json(json_data, output_filename="tailored_resume.docx"):
    doc = Document()

    title = doc.add_heading(json_data.get("name", ""), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph(
        " | ".join(filter(None, [
            json_data.get("phone", ""),
            json_data.get("email", ""),
            json_data.get("linkedin", ""),
            json_data.get("github", ""),
        ]))
    )
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(json_data.get("summary", ""))

    doc.add_heading("Top Skills", level=1)
    doc.add_paragraph(", ".join(json_data.get("skills", [])))

    doc.add_heading("Experience", level=1)
    for job in json_data.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{job.get('title', '')} — {job.get('company', '')}")
        run.bold = True
        p.add_run(f"  ({job.get('duration', '')})").italic = True
        for bullet in job.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Projects", level=1)
    for proj in json_data.get("projects", []):
        p = doc.add_paragraph()
        p.add_run(proj.get("name", "")).bold = True
        doc.add_paragraph(proj.get("description", ""))

    doc.save(output_filename)
    return output_filename
